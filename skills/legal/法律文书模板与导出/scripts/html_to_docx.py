#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert semantic legal-document HTML into DOCX using a JSON profile."""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SKILL_DIR = Path(__file__).resolve().parents[1]
PROFILES_DIR = SKILL_DIR / "assets" / "profiles"
DEFAULT_PROFILE = "fallback_desktop_word"
ALLOWED_PREFLIGHT_STATUS = {"PASS", "FIXED_PASS"}
MATTER_ROOT = Path(os.environ.get("LEGAL_WORKSPACE", ".")).expanduser()
SYSTEM_RECORD_ROOT = MATTER_ROOT / "_系统记录"
DRAFT_OUTPUT_NAME_PATTERN = re.compile(r"draft|unchecked|草稿|实验|未审查|未出稿")
INVALID_XML_CHARS_PATTERN = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


@dataclass
class Node:
    tag: str
    attrs: dict[str, str] = field(default_factory=dict)
    children: list["Node"] = field(default_factory=list)
    text_parts: list[str] = field(default_factory=list)


class TreeBuilder(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.root = Node("document")
        self.stack = [self.root]

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() in {"meta", "br", "hr", "img", "link", "input"}:
            return
        node = Node(tag.lower(), {k: (v or "") for k, v in attrs})
        self.stack[-1].children.append(node)
        self.stack.append(node)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        for idx in range(len(self.stack) - 1, 0, -1):
            if self.stack[idx].tag == tag:
                del self.stack[idx:]
                return

    def handle_data(self, data: str) -> None:
        if data:
            self.stack[-1].text_parts.append(data)


def sanitize_docx_text(text: str) -> str:
    return INVALID_XML_CHARS_PATTERN.sub("", text)


def text_of(node: Node) -> str:
    parts = list(node.text_parts)
    for child in node.children:
        parts.append(text_of(child))
    return sanitize_docx_text(" ".join("".join(parts).split()))


def first_descendant(node: Node, tag: str) -> Node | None:
    if node.tag == tag:
        return node
    for child in node.children:
        found = first_descendant(child, tag)
        if found:
            return found
    return None


def document_children(root: Node) -> list[Node]:
    body = first_descendant(root, "body") or root
    article = first_descendant(body, "article")
    return article.children if article else body.children


def has_nested_table(node: Node, in_table: bool = False) -> bool:
    if node.tag == "table" and in_table:
        return True
    child_in_table = in_table or node.tag == "table"
    return any(has_nested_table(child, child_in_table) for child in node.children)


def validate_html_tree(root: Node) -> None:
    title = first_descendant(root, "h1")
    if title is None or not text_of(title):
        raise ValueError("formal DOCX input requires non-empty h1 title")
    if has_nested_table(root):
        raise ValueError("nested table is not supported")


def load_profile(name: str | None) -> dict:
    profile_name = name or DEFAULT_PROFILE
    path = PROFILES_DIR / f"{profile_name}.json"
    if not path.exists():
        path = PROFILES_DIR / f"{DEFAULT_PROFILE}.json"
    with path.open("r", encoding="utf-8") as f:
        profile = json.load(f)
    profile["_path"] = str(path)
    return profile


def set_run_font(run, font_spec: dict) -> None:
    east_asia = font_spec.get("east_asia", "宋体")
    ascii_font = font_spec.get("ascii", "Times New Roman")
    run.font.name = east_asia
    run.font.size = Pt(float(font_spec.get("size_pt", 12)))
    run.font.bold = bool(font_spec.get("bold", False))
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def apply_paragraph_format(paragraph, profile: dict, *, first_indent: bool, alignment) -> None:
    para = profile["paragraph"]
    fmt = paragraph.paragraph_format
    paragraph.alignment = alignment
    fmt.space_before = Pt(float(para.get("space_before_pt", 0)))
    fmt.space_after = Pt(float(para.get("space_after_pt", 0)))
    fmt.first_line_indent = Pt(float(para.get("first_line_indent_pt", 0))) if first_indent else Pt(0)
    if para.get("line_spacing_rule") == "exact":
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(float(para.get("line_spacing_pt", 27)))
    else:
        fmt.line_spacing = float(para.get("line_spacing", 1.5))


def set_normal_style(doc: Document, profile: dict) -> None:
    normal = doc.styles["Normal"]
    body_font = profile["fonts"]["body"]
    normal.font.name = body_font.get("east_asia", "宋体")
    normal.font.size = Pt(float(body_font.get("size_pt", 12)))
    r_pr = normal._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), body_font.get("east_asia", "宋体"))
    r_fonts.set(qn("w:ascii"), body_font.get("ascii", "Times New Roman"))
    r_fonts.set(qn("w:hAnsi"), body_font.get("ascii", "Times New Roman"))


def add_page_number(paragraph, profile: dict) -> None:
    page_number = profile.get("page_number", {})
    if not page_number.get("enabled", True):
        return
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_font = profile["fonts"].get("footer", profile["fonts"]["body"])
    before = page_number.get("text_before", "")
    after = page_number.get("text_after", "")
    if before:
        run = paragraph.add_run(before)
        set_run_font(run, footer_font)
    for kind, text in [("begin", None), ("instr", "PAGE"), ("separate", None), ("text", "1"), ("end", None)]:
        run = paragraph.add_run()
        if kind == "instr":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = text
            run._r.append(node)
        elif kind == "text":
            node = OxmlElement("w:t")
            node.text = text
            run._r.append(node)
        else:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
            run._r.append(node)
    if after:
        run = paragraph.add_run(after)
        set_run_font(run, footer_font)


def set_cell_margins(cell, twip: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name in ["top", "start", "bottom", "end"]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(twip))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_width(table, width_pct: int = 5000) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(width_pct))


def table_rows(node: Node) -> list[list[tuple[str, str]]]:
    rows: list[list[tuple[str, str]]] = []
    for tr in descendants_by_tag(node, "tr"):
        row: list[tuple[str, str]] = []
        for cell in tr.children:
            if cell.tag in {"th", "td"}:
                row.append((cell.tag, text_of(cell)))
        if row:
            rows.append(row)
    return rows


def descendants_by_tag(node: Node, tag: str) -> Iterable[Node]:
    for child in node.children:
        if child.tag == tag:
            yield child
        yield from descendants_by_tag(child, tag)


def add_table(doc: Document, node: Node, profile: dict) -> None:
    rows = table_rows(node)
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    table_profile = profile.get("table", {})
    table_font = profile["fonts"].get("table", profile["fonts"]["body"])
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            tag, value = row[c_idx] if c_idx < len(row) else ("td", "")
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, int(table_profile.get("cell_margin_twip", 120)))
            if r_idx == 0 or tag == "th":
                set_cell_shading(cell, table_profile.get("header_fill", "F2F2F2"))
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if (r_idx == 0 or tag == "th" or c_idx == 0) else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(float(table_profile.get("line_spacing_pt", 18)))
            font_spec = dict(table_font)
            if r_idx == 0 or tag == "th":
                font_spec["bold"] = True
            run = paragraph.add_run(value)
            set_run_font(run, font_spec)


def add_text_paragraph(doc: Document, node: Node, profile: dict) -> None:
    text = text_of(node)
    if not text:
        return
    tag = node.tag
    classes = set(node.attrs.get("class", "").split())
    fonts = profile["fonts"]
    if tag == "h1":
        paragraph = doc.add_paragraph()
        apply_paragraph_format(paragraph, profile, first_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        paragraph.paragraph_format.space_after = Pt(12)
        font_spec = fonts["title"]
    elif tag == "h2":
        paragraph = doc.add_paragraph()
        apply_paragraph_format(paragraph, profile, first_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.paragraph_format.space_before = Pt(6)
        font_spec = fonts.get("heading1", fonts["body"])
    elif tag == "h3":
        paragraph = doc.add_paragraph()
        apply_paragraph_format(paragraph, profile, first_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.paragraph_format.space_before = Pt(4)
        font_spec = fonts.get("heading2", fonts["body"])
    elif "signature" in classes:
        paragraph = doc.add_paragraph()
        apply_paragraph_format(paragraph, profile, first_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        font_spec = fonts["body"]
    elif "meta" in classes or "label" in classes:
        paragraph = doc.add_paragraph()
        apply_paragraph_format(paragraph, profile, first_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        font_spec = fonts.get("meta", fonts["body"])
    elif "center" in classes or "center-note" in classes:
        paragraph = doc.add_paragraph()
        apply_paragraph_format(paragraph, profile, first_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        font_spec = fonts.get("meta", fonts["body"])
    else:
        paragraph = doc.add_paragraph()
        apply_paragraph_format(paragraph, profile, first_indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        font_spec = fonts["body"]
    run = paragraph.add_run(text)
    set_run_font(run, font_spec)


def render_nodes(doc: Document, nodes: Iterable[Node], profile: dict) -> None:
    for node in nodes:
        if node.tag in {"h1", "h2", "h3", "p"}:
            add_text_paragraph(doc, node, profile)
        elif node.tag == "table":
            add_table(doc, node, profile)
            doc.add_paragraph()
        elif node.tag in {"section", "article", "div", "body", "html"}:
            render_nodes(doc, node.children, profile)
        elif node.tag in {"ul", "ol"}:
            for li in [child for child in node.children if child.tag == "li"]:
                add_text_paragraph(doc, Node("p", {"class": "body"}, text_parts=[text_of(li)]), profile)
        elif text_of(node):
            add_text_paragraph(doc, Node("p", {"class": "body"}, text_parts=[text_of(node)]), profile)


def preflight_status(report_path: Path) -> str | None:
    if not report_path.exists():
        raise FileNotFoundError(f"preflight report not found: {report_path}")
    for line in report_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        if line.startswith("review_status:"):
            return line.split(":", 1)[1].strip()
    return None


def validate_preflight(input_path: Path, report_path: Path | None, allow_unchecked: bool) -> None:
    if allow_unchecked:
        return
    if input_path.name != "draft_checked.html":
        raise ValueError("formal DOCX export requires input file named draft_checked.html")
    if report_path is None:
        raise ValueError("formal DOCX export requires --preflight-report")
    status = preflight_status(report_path)
    if status not in ALLOWED_PREFLIGHT_STATUS:
        raise ValueError(f"preflight report status is not allowed: {status or 'missing'}")


def is_relative_to(path: Path, parent: Path) -> bool:
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except ValueError:
        return False


def validate_output_path(output_path: Path, allow_unchecked: bool) -> None:
    resolved = output_path.expanduser().resolve()
    if resolved.suffix.lower() != ".docx":
        raise ValueError("DOCX output path must end with .docx")

    if allow_unchecked:
        if is_relative_to(resolved, MATTER_ROOT):
            raise ValueError("unchecked DOCX export must not write into formal business area")
        return

    if ".cache" in resolved.parts:
        raise ValueError("formal DOCX output must not be under .cache")
    if is_relative_to(resolved, SYSTEM_RECORD_ROOT):
        raise ValueError("formal DOCX output must not be under system record root")
    if not is_relative_to(resolved, MATTER_ROOT):
        raise ValueError("formal DOCX output must be under business matter root")
    if DRAFT_OUTPUT_NAME_PATTERN.search(resolved.name.lower()):
        raise ValueError("formal DOCX filename must not look like draft or experiment")


def convert(input_path: Path, output_path: Path, profile_name: str | None) -> Path:
    profile = load_profile(profile_name)
    parser = TreeBuilder()
    parser.feed(input_path.read_text(encoding="utf-8"))
    validate_html_tree(parser.root)
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    page = profile["page"]
    section.page_width = Cm(float(page["width_cm"]))
    section.page_height = Cm(float(page["height_cm"]))
    if page.get("orientation") == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Cm(float(page["margin_top_cm"]))
    section.bottom_margin = Cm(float(page["margin_bottom_cm"]))
    section.left_margin = Cm(float(page["margin_left_cm"]))
    section.right_margin = Cm(float(page["margin_right_cm"]))
    if "header_distance_cm" in page:
        section.header_distance = Cm(float(page["header_distance_cm"]))
    if "footer_distance_cm" in page:
        section.footer_distance = Cm(float(page["footer_distance_cm"]))
    set_normal_style(doc, profile)
    add_page_number(section.footer.paragraphs[0], profile)
    render_nodes(doc, document_children(parser.root), profile)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def structural_check(path: Path) -> dict[str, bool]:
    with ZipFile(path) as zf:
        names = set(zf.namelist())
        document_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        footer_xml = "".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.startswith("word/footer") and name.endswith(".xml")
        )
    return {
        "document_xml": "word/document.xml" in names,
        "page_margin": "<w:pgMar" in document_xml,
        "page_number": "PAGE" in footer_xml,
        "real_table": "<w:tbl>" in document_xml,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert semantic legal HTML to DOCX.")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--profile", default=DEFAULT_PROFILE)
    parser.add_argument("--preflight-report", type=Path)
    parser.add_argument("--allow-unchecked", action="store_true", help="Only for explicit non-deliverable experiments.")
    parser.add_argument("--check", action="store_true", help="Print structural checks after export.")
    args = parser.parse_args()
    validate_preflight(args.input, args.preflight_report, args.allow_unchecked)
    validate_output_path(args.output, args.allow_unchecked)
    out = convert(args.input, args.output, args.profile)
    print(f"DOCX generated: {out}")
    if args.check:
        for key, ok in structural_check(out).items():
            print(f"{key}: {ok}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
